import tempfile
from pathlib import Path

import fitz
import pytest

from backend.app.parser import ParseError, parse_document

SAMPLES_DIR = Path("data/samples")


class TestParsePDF:
    def test_extracts_text_from_real_pdf(self):
        pdf_path = SAMPLES_DIR / "electrical_spec_luch" / "document.pdf"
        raw_text, sections = parse_document(pdf_path)
        assert len(raw_text) > 1000
        assert len(sections) > 0

    def test_sections_have_valid_pages(self):
        pdf_path = SAMPLES_DIR / "electrical_spec_luch" / "document.pdf"
        _, sections = parse_document(pdf_path)
        for s in sections:
            assert s.page_start >= 1
            assert s.page_end >= s.page_start
            assert len(s.name) > 0
            assert len(s.text) > 0

    def test_ventilation_pdf(self):
        pdf_path = SAMPLES_DIR / "ventilation_fosfokhim" / "document.pdf"
        raw_text, sections = parse_document(pdf_path)
        assert len(raw_text) > 1000
        assert len(sections) > 0


class TestParseErrors:
    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            parse_document("/tmp/nonexistent_file.pdf")

    def test_unsupported_format(self):
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            f.write(b"test")
            tmp = f.name
        try:
            with pytest.raises(ParseError, match="Неподдерживаемый формат"):
                parse_document(tmp)
        finally:
            Path(tmp).unlink()

    def test_empty_pdf_no_text_layer(self):
        """PDF без текстового слоя должен вызвать ParseError."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            tmp = f.name
        doc = fitz.open()
        doc.new_page()
        doc.save(tmp)
        doc.close()
        try:
            with pytest.raises(ParseError, match="текстового слоя"):
                parse_document(tmp)
        finally:
            Path(tmp).unlink()


class TestParseDOCX:
    def test_parse_docx(self):
        """Тест парсинга простого DOCX-файла."""
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp = f.name

        doc = Document()
        doc.add_heading("1. Общие сведения", level=1)
        doc.add_paragraph("Проект разработан для объекта строительства. " * 20)
        doc.add_heading("2. Электроснабжение", level=1)
        doc.add_paragraph("Электроснабжение выполняется от трансформаторной подстанции. " * 20)
        doc.save(tmp)

        try:
            raw_text, sections = parse_document(tmp)
            assert len(raw_text) > MIN_TEXT_THRESHOLD
            assert len(sections) >= 1
        finally:
            Path(tmp).unlink()


# Минимальный порог длины текста из parser.py
MIN_TEXT_THRESHOLD = 200
